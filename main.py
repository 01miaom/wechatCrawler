from selenium import webdriver
from selenium.webdriver.support.ui import Select
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
import requests
import numpy as np
import pandas as pd
import time
from urllib import request
from urllib.request import Request,urlopen
import json
from jsonpath import jsonpath
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches
import os
from urllib.request import urlretrieve

browser=webdriver.Chrome(chromedriver")

N = len(d)# d is a .csv file with 2 columns: article 'name' and 'url'
                         
for j in range(N):
    namep = d['name'][j]
    urlp = d['url'][j]
    np = j
    browser.get(urlp)
    nimg = len(browser.find_elements_by_xpath('//img'))
    filename = './公众号/'+str(j)+"."+namep+'.docx'
    # save the image in a folder
    for i in range(nimg):
        imgurl = browser.find_elements_by_xpath('//img')[i].get_attribute('src')
        path = './公众号/img/p'+str(np)+'img'+str(i)+'.png'
        try:
            urlretrieve(imgurl,path)
        except:
            print("can not get the img")
        time.sleep(1)
    p = browser.find_elements_by_class_name('rich_media_wrp')[0].text
    document = Document()
    document.styles['Normal'].font.name=u'宋体'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'),u'宋体')
    # write the article and iamge to a .docx file
    document.add_paragraph(p)
    for i in range(nimg):
        path = './公众号/img/p'+str(np)+'img'+str(i)+'.png'
        try:
            document.add_picture(path, width=Inches(5))
        except:
            print("no file")
    document.save(filename)
    time.sleep(5)

